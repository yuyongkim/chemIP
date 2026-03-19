"""Export submission report from Markdown to TXT/RTF/DOCX.

Usage:
  python scripts/export_submission_report.py
  python scripts/export_submission_report.py --input docs/SUBMISSION_REPORT.md
"""

from __future__ import annotations

import argparse
import re
import sys
import struct
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


def configure_output() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        try:
            sys.stdout.reconfigure(encoding="utf-8", errors="replace")
            sys.stderr.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass


def normalize_inline(text: str) -> str:
    value = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"[IMAGE] \1 (\2)", text)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"__([^_]+)__", r"\1", value)
    value = re.sub(r"\*([^*]+)\*", r"\1", value)
    return value.strip()


def strip_markdown(md: str) -> str:
    text = md.replace("\r\n", "\n").replace("\r", "\n")

    # Preserve code block content, remove fences.
    text = re.sub(r"```[a-zA-Z0-9_-]*\n", "", text)
    text = text.replace("```", "")

    # Images and links.
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"[IMAGE] \1 (\2)", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)

    # Headers.
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)

    # Tables: drop separator lines, keep rows as tab-separated text.
    lines: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        if re.match(r"^\|?[\s:-]+\|[\s|:-]*$", stripped):
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            line = "\t".join(cells)
        lines.append(line)
    text = "\n".join(lines)

    # Inline formatting.
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)

    # Normalize extra blank lines.
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text + "\n"


def text_to_rtf(text: str) -> str:
    def encode_char(ch: str) -> str:
        code = ord(ch)
        if ch == "\\":
            return r"\\"
        if ch == "{":
            return r"\{"
        if ch == "}":
            return r"\}"
        if ch == "\n":
            return r"\par" + "\n"
        if code < 128:
            return ch
        signed = code if code <= 32767 else code - 65536
        return rf"\u{signed}?"

    body = "".join(encode_char(ch) for ch in text)
    return "{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0 Calibri;}}\n\\f0\\fs22\n" + body + "\n}"


def is_table_separator_line(line: str) -> bool:
    return bool(re.match(r"^\|?[\s:-]+\|[\s|:-]*$", line.strip()))


def add_table_from_lines(document: Document, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        if is_table_separator_line(stripped):
            continue
        cells = [normalize_inline(cell.strip()) for cell in stripped.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value

    document.add_paragraph("")


def try_add_image(document: Document, line: str, base_dir: Path) -> bool:
    match = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
    if not match:
        return False

    alt_text = normalize_inline(match.group(1))
    path_text = match.group(2).strip()
    image_path = resolve_image_path(base_dir, path_text)

    if image_path and image_path.exists():
        try:
            width, height = get_image_size(image_path)
            if width and height:
                # Keep image within a single page even with surrounding heading text.
                max_width_in = 5.8
                max_height_in = 8.0
                ratio = height / width
                target_height = max_width_in * ratio
                if target_height <= max_height_in:
                    document.add_picture(str(image_path), width=Inches(max_width_in))
                else:
                    document.add_picture(str(image_path), height=Inches(max_height_in))
            else:
                document.add_picture(str(image_path), width=Inches(5.8))
            return True
        except Exception:
            pass

    document.add_paragraph(f"[IMAGE] {alt_text} ({path_text})")
    document.add_paragraph("")
    return True


def resolve_image_path(base_dir: Path, path_text: str) -> Path | None:
    raw = Path(path_text)
    candidates: list[Path] = []

    if raw.is_absolute():
        candidates.append(raw)
    else:
        candidates.append(base_dir / raw)
        candidates.append(base_dir.parent / raw)
        candidates.append(base_dir.parent.parent / raw)

    for candidate in candidates:
        resolved = candidate.resolve()
        if resolved.exists():
            return resolved
    return candidates[0].resolve() if candidates else None


def get_image_size(image_path: Path) -> tuple[int | None, int | None]:
    suffix = image_path.suffix.lower()
    if suffix == ".png":
        try:
            with image_path.open("rb") as f:
                header = f.read(24)
            if len(header) >= 24 and header[:8] == b"\x89PNG\r\n\x1a\n":
                width, height = struct.unpack(">II", header[16:24])
                return width, height
        except Exception:
            return None, None
    return None, None


def set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_style_font(document: Document, style_name: str, font_name: str) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return

    style.font.name = font_name
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def apply_korean_font(document: Document, font_name: str = "맑은 고딕") -> None:
    style_names = [
        "Normal",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "List Bullet",
        "List Number",
    ]
    for style_name in style_names:
        set_style_font(document, style_name, font_name)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, font_name)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, font_name)


def markdown_to_docx(md: str, output_path: Path, base_dir: Path) -> None:
    document = Document()
    lines = md.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    in_code = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            document.add_paragraph(line)
            i += 1
            continue

        if not stripped:
            document.add_paragraph("")
            i += 1
            continue

        if try_add_image(document, line, base_dir):
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            block: list[str] = []
            while i < len(lines):
                probe = lines[i].strip()
                if probe.startswith("|") and probe.endswith("|"):
                    block.append(lines[i])
                    i += 1
                    continue
                break
            add_table_from_lines(document, block)
            continue

        heading = re.match(r"^\s{0,3}(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = normalize_inline(heading.group(2))
            document.add_heading(text, level=min(level, 4))
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            # Preserve explicit numbering text from markdown to avoid global auto-number continuation.
            document.add_paragraph(normalize_inline(line))
            i += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            document.add_paragraph(normalize_inline(bullet.group(1)), style="List Bullet")
            i += 1
            continue

        document.add_paragraph(normalize_inline(line))
        i += 1

    apply_korean_font(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def resolve_writable_path(path: Path) -> Path:
    if not path.exists():
        return path

    try:
        with path.open("a", encoding="utf-8"):
            return path
    except PermissionError:
        pass

    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    candidate = parent / f"{stem}_new{suffix}"
    index = 2
    while candidate.exists():
        candidate = parent / f"{stem}_new{index}{suffix}"
        index += 1
    return candidate


def export_files(input_path: Path) -> tuple[Path, Path, Path]:
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    md = input_path.read_text(encoding="utf-8", errors="replace")
    txt = strip_markdown(md)
    rtf = text_to_rtf(txt)

    txt_path = resolve_writable_path(input_path.with_suffix(".txt"))
    rtf_path = resolve_writable_path(input_path.with_suffix(".rtf"))
    docx_path = resolve_writable_path(input_path.with_suffix(".docx"))

    txt_path.write_text(txt, encoding="utf-8")
    rtf_path.write_text(rtf, encoding="utf-8")
    markdown_to_docx(md, docx_path, input_path.parent)
    return txt_path, rtf_path, docx_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export submission report to TXT/RTF/DOCX.")
    parser.add_argument(
        "--input",
        default="docs/SUBMISSION_REPORT.md",
        help="Input markdown file path (default: docs/SUBMISSION_REPORT.md)",
    )
    return parser.parse_args()


def main() -> int:
    configure_output()
    args = parse_args()
    input_path = Path(args.input)
    try:
        txt_path, rtf_path, docx_path = export_files(input_path)
        print(f"[PASS] TXT exported: {txt_path}")
        print(f"[PASS] RTF exported: {rtf_path}")
        print(f"[PASS] DOCX exported: {docx_path}")
        return 0
    except Exception as exc:
        print(f"[FAIL] {exc}")
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
