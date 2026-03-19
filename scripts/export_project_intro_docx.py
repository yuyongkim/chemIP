"""Export project intro markdown to a Word (.docx) document.

Usage:
  .venv\\Scripts\\python.exe scripts/export_project_intro_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT_MD = ROOT / "docs" / "CHEMIP_PROJECT_INTRO_AND_SCREENSHOT_GUIDE_2026-03-02.md"
OUTPUT_DOCX = ROOT / "docs" / "CHEMIP_PROJECT_INTRO_AND_SCREENSHOT_GUIDE_2026-03-02.docx"
SCREENSHOT_DIR = ROOT / "docs" / "screenshots" / "proposal_20260302"

INLINE_FIGURES = {
    "1.3 핵심 가치 제안": ("스크린샷 #6 결과/성과 화면", SCREENSHOT_DIR / "06_result_ai_insight.png"),
    "2.1 주요 기능 설명": ("스크린샷 #2 핵심 기능 화면", SCREENSHOT_DIR / "02_core_chemical_detail.png"),
    "2.2 사용자 여정": ("스크린샷 #3 대시보드/관리 화면", SCREENSHOT_DIR / "03_dashboard_trade.png"),
    "2.3 화면 구성": ("스크린샷 #1 메인 화면/랜딩페이지", SCREENSHOT_DIR / "01_landing_main.png"),
    "3.1 기술 아키텍처": ("스크린샷 #5 모바일/반응형 화면", SCREENSHOT_DIR / "05_mobile_home.png"),
    "3.4 보안 및 데이터 보호": ("스크린샷 #4 설정/환경설정 화면", SCREENSHOT_DIR / "04_settings_ops_guide.png"),
}


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table_row(line: str) -> list[str]:
    parts = [part.strip() for part in line.strip().strip("|").split("|")]
    return parts


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    for raw in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(raw.rstrip("\n"))
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return

    rows = [parse_table_row(line) for line in table_lines]
    header = rows[0]

    # Skip markdown separator row like | --- | --- |
    body = rows[1:]
    if body and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in body[0]):
        body = body[1:]

    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light List Accent 1"

    for idx, value in enumerate(header):
        table.rows[0].cells[idx].text = value

    for r_idx, row in enumerate(body, start=1):
        for c_idx, value in enumerate(row):
            if c_idx < len(table.rows[r_idx].cells):
                table.rows[r_idx].cells[c_idx].text = value

    doc.add_paragraph("")


def add_inline_figure(doc: Document, section_title: str) -> None:
    figure = INLINE_FIGURES.get(section_title)
    if not figure:
        return

    caption, image_path = figure
    doc.add_paragraph(caption)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.0))
        doc.add_paragraph(str(image_path.relative_to(ROOT)))
    else:
        doc.add_paragraph(f"[MISSING IMAGE] {image_path}")
    doc.add_paragraph("")


def add_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
                doc.add_paragraph("")
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table block
        if is_table_line(line):
            table_lines = [line]
            i += 1
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            doc.add_heading(title, level=1)
            add_inline_figure(doc, title)
            i += 1
            continue
        if stripped.startswith("## "):
            title = stripped[3:].strip()
            doc.add_heading(title, level=2)
            add_inline_figure(doc, title)
            i += 1
            continue
        if stripped.startswith("### "):
            title = stripped[4:].strip()
            doc.add_heading(title, level=3)
            add_inline_figure(doc, title)
            i += 1
            continue
        if stripped.startswith("#### "):
            title = stripped[5:].strip()
            doc.add_heading(title, level=4)
            add_inline_figure(doc, title)
            i += 1
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            body = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(body, style="List Number")
            i += 1
            continue

        if stripped == "":
            doc.add_paragraph("")
            i += 1
            continue

        doc.add_paragraph(line)
        i += 1


def main() -> int:
    if not INPUT_MD.exists():
        raise FileNotFoundError(f"Input markdown not found: {INPUT_MD}")

    md_text = INPUT_MD.read_text(encoding="utf-8")

    doc = Document()
    set_default_font(doc)
    add_markdown(doc, md_text)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(f"[OK] {OUTPUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
