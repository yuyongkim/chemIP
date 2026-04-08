#!/usr/bin/env python3
"""Build a properly formatted ACS-style Word document from paper markdown."""
import sys
import io
import re
import os

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


FIGURE_MAP = {
    1: ("papers/acs-chs/figures/fig1_architecture.png",
        "Figure 1. Three-tier architecture of the ChemIP platform showing the browser-based presentation layer, FastAPI server with adapter modules, and external data sources."),
    2: ("papers/acs-chs/figures/cap_main.png",
        "Figure 2. ChemIP main interface with tabbed navigation for safety, regulatory, patent, pharmaceutical, and AI synthesis workflows."),
    3: ("papers/acs-chs/figures/cap_chemical.png",
        "Figure 3. Substance safety lookup showing MSDS data with GHS classifications, hazard statements, and precautionary information."),
    4: ("papers/acs-chs/figures/cap_patents.png",
        "Figure 4. Patent landscape search interface combining KIPRIS API results with optional local index lookups."),
    5: ("papers/acs-chs/figures/cap_drugs.png",
        "Figure 5. Pharmaceutical cross-reference aggregating Korean drug approvals (MFDS), US drug labels (OpenFDA), and biomedical literature (PubMed)."),
    6: ("papers/acs-chs/figures/fig6_workflow_comparison.png",
        "Figure 6. Comparison of conventional multi-portal workflow (left, ~150 min) and ChemIP single-query workflow (right, ~35 min) for per-substance safety evaluation."),
}

FIGURES_INSERTED = set()


def insert_figure(doc, fig_num):
    """Insert a figure image with caption after the current paragraph."""
    if fig_num in FIGURES_INSERTED:
        return
    if fig_num not in FIGURE_MAP:
        return
    img_path, caption = FIGURE_MAP[fig_num]
    if not os.path.exists(img_path):
        return

    # Add image
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(4)
    run = p_img.add_run()
    run.add_picture(img_path, width=Inches(5.5))

    # Add caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run = p_cap.add_run(caption)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True

    FIGURES_INSERTED.add(fig_num)


def add_rich_text(paragraph, text):
    """Parse markdown inline formatting and add runs with proper formatting."""
    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*|\^(?:[\da-z]+[,-]?)+)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("^") and re.match(r"^\^[\da-z]+([,-][\da-z]+)*$", part):
            run = paragraph.add_run(part[1:])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def set_double_spacing(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def parse_md_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        cells = [c.strip() for c in line.split("|")]
        # Skip separator rows (lines like |---|:---:|-----|)
        is_separator = all(re.match(r"^:?-+:?$", c.strip()) for c in cells if c.strip())
        if is_separator:
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows, caption=None):
    """Add a properly formatted Word table."""
    if caption:
        p = doc.add_paragraph()
        set_double_spacing(p)
        add_rich_text(p, caption)
        if p.runs:
            # Make "Table N." bold
            pass

    if not rows:
        return

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                add_rich_text(p, cell_text)
                if i == 0:
                    for run in p.runs:
                        run.bold = True
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"


def add_page_numbers(doc):
    """Add page numbers to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fldChar1 = parse_xml(
        '<w:fldChar %s w:fldCharType="begin"/>' % nsdecls("w")
    )
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        '<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls("w")
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(
        '<w:fldChar %s w:fldCharType="end"/>' % nsdecls("w")
    )
    run3._r.append(fldChar2)


def build_docx(md_path, out_path):
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Strip YAML frontmatter
    content = re.sub(r"^---.*?---\s*", "", content, flags=re.DOTALL)

    doc = Document()

    # ===== STYLES =====
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    for level in range(1, 4):
        hs = doc.styles["Heading %d" % level]
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(14 if level == 1 else 12)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    # ===== PAGE SETUP =====
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_page_numbers(doc)

    # ===== PARSE AND BUILD =====
    lines = content.split("\n")
    i = 0
    table_buffer = []
    table_caption = None
    in_table = False
    title_added = False

    while i < len(lines):
        line = lines[i]

        # Heading
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            # Flush table
            if table_buffer:
                rows = parse_md_table(table_buffer)
                add_table(doc, rows, table_caption)
                table_buffer = []
                table_caption = None
                in_table = False

            level = len(heading_match.group(1))
            text = heading_match.group(2)

            if level == 1 and not title_added:
                # Title
                p = doc.add_paragraph(text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
                    run.bold = True

                # Author block
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_double_spacing(p2)
                run = p2.add_run("Yuyong Kim")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

                p3 = doc.add_paragraph()
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_double_spacing(p3)
                run = p3.add_run("University of Wisconsin-Madison, United States")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

                p4 = doc.add_paragraph()
                p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_double_spacing(p4)
                run = p4.add_run("Email: ykim288@wisc.edu")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

                p5 = doc.add_paragraph()
                p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_double_spacing(p5)
                run = p5.add_run("ORCID: 0009-0006-4842-666X")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

                title_added = True
                i += 1
                continue

            p = doc.add_heading(text, level=level)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = "Times New Roman"
            i += 1
            continue

        # Table line
        if "|" in line and line.strip().startswith("|") and line.strip().count("|") >= 3:
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            rows = parse_md_table(table_buffer)
            add_table(doc, rows, table_caption)
            table_buffer = []
            table_caption = None
            in_table = False

        # Table caption
        stripped = line.strip()
        if re.match(r"^\*?\*?Table \d+\.", stripped):
            table_caption = stripped
            i += 1
            continue

        # Footnotes (^a, ^b) - may have multiple on one line
        if stripped.startswith("^") and re.match(r"^\^[a-z]", stripped):
            p = doc.add_paragraph()
            set_double_spacing(p)
            # Split on footnote markers
            parts = re.split(r"(\^[a-z])\s*", stripped)
            for part in parts:
                if not part:
                    continue
                if re.match(r"^\^[a-z]$", part):
                    run = p.add_run(part[1:])
                    run.font.superscript = True
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"
                else:
                    run = p.add_run(" " + part.strip() + " ")
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
            i += 1
            continue

        # Numbered list (references section)
        ref_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ref_match:
            p = doc.add_paragraph()
            set_double_spacing(p)
            num = ref_match.group(1)
            text = ref_match.group(2)
            run_num = p.add_run("(%s) " % num)
            run_num.font.name = "Times New Roman"
            add_rich_text(p, text)
            for run in p.runs:
                if run.font.name is None:
                    run.font.name = "Times New Roman"

            # Check for figure references
            fig_refs = re.findall(r"\(Figure (\d+)\)", stripped)
            for ref in fig_refs:
                insert_figure(doc, int(ref))

            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        set_double_spacing(p)
        add_rich_text(p, line)
        for run in p.runs:
            if run.font.name is None:
                run.font.name = "Times New Roman"

        # Check for figure references and insert after paragraph
        fig_refs = re.findall(r"\(Figure (\d+)\)", line)
        for ref in fig_refs:
            insert_figure(doc, int(ref))

        i += 1

    # Flush remaining table
    if table_buffer:
        rows = parse_md_table(table_buffer)
        add_table(doc, rows, table_caption)

    doc.save(out_path)

    # Verify
    doc2 = Document(out_path)
    bold_c = sum(1 for p in doc2.paragraphs for r in p.runs if r.bold)
    italic_c = sum(1 for p in doc2.paragraphs for r in p.runs if r.italic)
    super_c = sum(1 for p in doc2.paragraphs for r in p.runs if r.font.superscript)
    print("Bold runs: %d, Italic runs: %d, Superscript runs: %d" % (bold_c, italic_c, super_c))
    print("Tables: %d" % len(doc2.tables))
    print("Paragraphs: %d" % len(doc2.paragraphs))
    print("DONE: %s" % out_path)


if __name__ == "__main__":
    build_docx("papers/acs-chs/paper_v1.7.md", "papers/acs-chs/paper_v1.7.docx")
